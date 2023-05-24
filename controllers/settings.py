import hashlib
import os
import datetime
import json
import time
import threading
from models import *
from flask import Flask, redirect, url_for, render_template, request, session, flash, jsonify, send_file, send_from_directory
from flask_session import Session  # https://pythonhosted.org/Flask-Session
from mongoengine.queryset.visitor import Q
from datetime import timedelta, datetime
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm

def settings():
    if not checkSession():
        sessionAuthUri = session["flow"]["auth_uri"]
        return redirect(sessionAuthUri)
    if request.method == "GET":
        user = User.objects(username=session["user"].get("name")).first()
        sessionUser = session["user"].get("name").replace(" ", "").lower()
        signatureFile = 1
        if not os.path.exists(f'{sessionUser}/signature.png'):
            signatureFile = 0
        warehouses = Warehouse.objects().all()
        return render_template("settings.html", user=user, warehouses=warehouses, signatureFile=signatureFile)


def changeSettings():
    if not checkSession():
        sessionAuthUri = session["flow"]["auth_uri"]
        return redirect(sessionAuthUri)
    else:
        sessionUser = session["user"].get("name").replace(" ", "").lower()
        locationAllocation = 0
        manualInputFields = 1
        useSignature = 0
        theme = 1
        signatureOffset = 0
        receivedLocation = request.form["receivedLocation"]

        signatureFile = request.files['signatureFile']
        if signatureFile and isImageFile(signatureFile.filename):
            filename = secure_filename(signatureFile.filename)
            extension = filename.rsplit('.', 1)[1].lower()
            signatureFile.save(f"{sessionUser}/signature.{extension}")
            if not extension == "png":
                im = Image.open(f"{sessionUser}/signature.{extension}")
                rgb_im = im.convert('RGB')
                rgb_im.save(f"{sessionUser}/signature.png")
            signatureOffset = makeImageTransparent()


        if request.form.get('locationAllocation'):
            locationAllocation = 1
        if request.form.get('useSignature'):
            useSignature = 1

        #print(theme)
        #print(receivedLocation)
        #print(locationAllocation)
        #print(manualInputFields)

        User.objects(username=session["user"].get("name")).update(
            set__settings__receivedLocation=receivedLocation, set__settings__theme=theme, set__settings__locationAllocation=locationAllocation, set__settings__manualInputFields=manualInputFields, set__settings__useSignature=useSignature, set__settings__signatureOffset=signatureOffset)

        return redirect(url_for("home"))


def signatureExample():
    sessionUser = session["user"].get("name").replace(" ", "").lower()
    if not os.path.exists(f'{sessionUser}/signature.png'):
        return "Signature has not been saved, thus none could be found."

    counter = 1
    document = Document(f"{sessionUser}/reverzTemplateBoth.docx")

    table = document.tables[1]
    table.cell(0, 0).paragraphs[0].add_run("\n" + "Uporabnik")
    table.cell(0, 1).paragraphs[0].add_run("\n" + "12345678")
    table.cell(0, 0).paragraphs[0].runs[0].font.bold = True

    table = document.tables[0]
    table.cell(0, 0).paragraphs[0].add_run(" " + "Odelek")
    table.cell(1, 0).paragraphs[0].add_run(" " + "12345678")

    user = session["user"].get("name")
    mongoMe = User.objects(username=user).first()
    employee = "Uporabnik"
    givenLocation = "Lokacija"
    receivedLocation = "Lokacija"
    reverzCounter = "0"
    signatureOffset = mongoMe["settings"]["signatureOffset"]
    star = ""

    document.save(f"{sessionUser}/tmp2.docx")
    os.system(
        f"./reverzComplete.sh '{sessionUser}' '{user}' '{employee}' '{givenLocation}' '{receivedLocation}' '{reverzCounter}' '{star}'")
    os.system(f"unoconv -f pdf {sessionUser}/tmp2.docx")
    
    print(signatureOffset)
    in_pdf_file = f'{sessionUser}/tmp2.pdf'
    out_pdf_file = f'{sessionUser}/tmp.pdf'
    img_file = f'{sessionUser}/signature.png'
    packet = io.BytesIO()
    can = canvas.Canvas(packet)
    x_start = 70
    y_start = 60 - int(signatureOffset) * 0.44 + 3
    can.drawImage(img_file, x_start, y_start, width=120, height=45, preserveAspectRatio=True, mask='auto')
    can.save()
    #move to the beginning of the StringIO buffer
    packet.seek(0)
    new_pdf = PdfFileReader(packet)
    # read the existing PDF
    existing_pdf = PdfFileReader(open(in_pdf_file, "rb"))
    output = PdfFileWriter()
    for i in range(len(existing_pdf.pages)):
        page = existing_pdf.getPage(i)
        page.mergePage(new_pdf.getPage(i))
        output.addPage(page)
    outputStream = open(out_pdf_file, "wb")
    output.write(outputStream)
    outputStream.close()
    return send_file(f"{sessionUser}/tmp.pdf", download_name=f"signature_example.pdf", as_attachment=True)
