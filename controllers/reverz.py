import hashlib
import os
import datetime
import json
import time
import threading
from models import *
from flask import Flask, redirect, url_for, render_template, request, session, flash, jsonify, send_file, send_from_directory
from flask_session import Session  # https://pythonhosted.org/Flask-Session
from mongoengine.queryset.visitor import Q
from datetime import timedelta, datetime
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm

def deleteReverz():

    if not checkSession():
        sessionAuthUri = session["flow"]["auth_uri"]
        return redirect(sessionAuthUri)

    reverzJson = request.get_json()
    reverzId = reverzJson["id"]
    goodbye = Reverz.objects(id=reverzId).first()
    goodbye2 = ReverzEquipment.objects(reverzId=reverzId).all()
    goodbye.delete()
    goodbye2.delete()
    os.system(f"rm log/{reverzId}.pdf")
    return jsonify("ok")


def reverz():
    if not checkSession():
        sessionAuthUri = session["flow"]["auth_uri"]
        return redirect(sessionAuthUri)
    if request.method == "POST":
        reverz = request.get_json()

        # Save json to database
        givenEq = []
        counter = 0
        newEquipment = []
        allEquipment = []

        mongoMe = User.objects(username=session["user"].get("name")).first()
        star = False

        for e in reverz["izdanaOprema"]:
            mongoEq = Equipment.objects(name=e).first()
            currentInvNum = reverz["izdanaInvNum"][counter]
            if currentInvNum[-1] == "*":
                star = True
                newEquipment.append(
                    [mongoMe["username"], reverz["uporabnik"], currentInvNum[:-1], e, str(datetime.now().strftime("%d. %m. %Y"))])
                currentInvNum = currentInvNum[:-1]

            givenEq.append(ImdbEquipment(
                equipment_id=mongoEq["id"], invNum=currentInvNum))
            counter = counter + 1
        receivedEq = []
        counter = 0
        for e in reverz["vrnjenaOprema"]:
            mongoEq = Equipment.objects(name=e).first()
            receivedEq.append(ImdbEquipment(
                equipment_id=mongoEq["id"], invNum=reverz["vrnjenaInvNum"][counter]))
            counter = counter + 1

        mongoReceiver = Employee.objects(name=reverz["uporabnik"]).first()
        mongoReverz = Reverz(giver=mongoMe["id"], receiver=mongoReceiver["id"], givenEq=givenEq,
                             receivedEq=receivedEq, givenLoc=reverz["izdanaLokacija"], receivedLoc=reverz["vrnjenaLokacija"], reverzCount=(Reverz.objects.count() + 100), date=datetime.now())
        Reverz.objects.insert(mongoReverz)

        # Shrani opremo posebej
        #print(mongoReverz["id"])
        for ge in mongoReverz["givenEq"]:
            reverzEquipment = ReverzEquipment(giver=mongoMe["username"], receiver=reverz["uporabnik"], equipmentInvNum=ge["invNum"],
                                              equipmentName=ge["equipment_id"]["name"], status="A", loc=reverz["izdanaLokacija"], date=mongoReverz["date"], reverzId=str(mongoReverz["id"]))
            reverzEquipment.save()
        for ree in mongoReverz["receivedEq"]:
            reverzEquipment = ReverzEquipment(giver=mongoMe["username"], receiver=reverz["uporabnik"], equipmentInvNum=ree["invNum"],
                                              equipmentName=ree["equipment_id"]["name"], status="S", loc=reverz["vrnjenaLokacija"], date=mongoReverz["date"], reverzId=str(mongoReverz["id"]))
            reverzEquipment.save()

        # Create and show reverz in pdf
        mongoUnit = Unit.objects(number=mongoReceiver["strm"]).first()
        reverz.update({"EZSO": str(mongoReceiver["ezso"])})
        reverz.update({"STRM": str(mongoUnit["number"])})
        reverz.update({"unit": str(mongoUnit["name"])})
        reverz.update({"me": str(mongoMe["username"])})
        reverz.update({"useSignature": str(mongoMe["settings"]["useSignature"])})
        reverz.update({"signatureOffset": str(mongoMe["settings"]["signatureOffset"])})
        reverz.update({"reverzCounter": str(mongoReverz["reverzCount"])})
        if star:
            reverz.update({"star": "1"})
        else:
            reverz.update({"star": "0"})
        sessionUser = session["user"].get("name").replace(" ", ".").lower()
        if newEquipment:
            threading.Thread(target=sendEmail, args=(newEquipment, reverz["izdanaLokacija"], mongoUnit["number"], mongoReceiver["ezso"], sessionUser)).start()
        mongoReceiver = Employee.objects(name=reverz["uporabnik"]).update(
            set__location=reverz["izdanaLokacija"])
        print(reverz["vrnjenaLokacija"])
        createReverz(reverz)

        sessionUser = session["user"].get("name").replace(" ", "").lower()
        os.system(f'cp {sessionUser}/tmp.pdf log/{mongoReverz["id"]}.pdf')

        receiver = reverz["uporabnik"]
        #reverz_date = str(mongoReverz["date"]).replace(":", "").replace(".","").replace(" ", "_").replace("-", "")
        reverz_date = mongoReverz["date"].strftime("%d-%m-%Y_%H.%M_") + str(mongoReverz["reverzCount"])

        os.system(f'cp {sessionUser}/tmp.pdf /media/doc/"{receiver}_{reverz_date}.pdf"')
        #os.system(f'cp {sessionUser}/tmp.pdf media/"{receiver}_{reverz_date}.pdf"')
        #send_file(f"{sessionUser}/tmp.pdf", download_name='reverz.pdf')

        return jsonify(str(mongoReverz["id"]), 200)

    if request.method == "GET":
        employees = Employee.objects().order_by("name")
        equipment = Equipment.objects().order_by("name")
        me = User.objects(username=session["user"].get("name")).first()
        producers = Producer.objects().all()
        equipmentTypes = EquipmentType.objects().all()
        units = Unit.objects().all()
        warehouses = Warehouse.objects().all()
        locationAllocation = me["settings"]["locationAllocation"]
        receivedLocation = me["settings"]["receivedLocation"]

        return render_template("reverzForm.html", employees=employees, equipment=equipment, sessionUser=session["user"].get("name").replace(" ", "").lower(), locationAllocation=locationAllocation, receivedLocation=receivedLocation, warehouses=warehouses, producers=producers, equipmentTypes=equipmentTypes, units=units, user=me)


def createReverz(reverzJson):
    sessionUser = session["user"].get("name").replace(" ", "").lower()
    both = False
    given = False
    if len(reverzJson["izdanaOprema"]) > 0 and len(reverzJson["vrnjenaOprema"]) > 0:
        both = True
    elif len(reverzJson["izdanaOprema"]) > 0:
        given = True

    counter = 1
    document = None
    if both:
        document = Document(f"{sessionUser}/reverzTemplateBoth.docx")
    elif given:
        document = Document(f"{sessionUser}/reverzTemplateIzdana.docx")
    else:
        document = Document(f"{sessionUser}/reverzTemplateVrnjena.docx")

    #print(document._body)
    if given or both:
        table = document.tables[2]
        for e in reverzJson["izdanaOprema"]:
            table.add_row()  # ADD ROW HERE
            table.cell(counter, 0).text = str(counter) + "."
            table.cell(counter, 1).text = e
            table.cell(counter, 2).text = reverzJson["izdanaInvNum"][counter-1]
            table.cell(
                counter, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(
                counter, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(
                counter, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.rows[counter].height = Cm(0.8)
            counter = counter + 1

    whichTable = 3
    counter = 1
    if not both and not given:
        whichTable = 2

    if not given:
        table = document.tables[whichTable]
        for e in reverzJson["vrnjenaOprema"]:
            table.add_row()  # ADD ROW HERE
            table.cell(counter, 0).text = str(counter) + "."
            table.cell(counter, 1).text = e
            table.cell(
                counter, 2).text = reverzJson["vrnjenaInvNum"][counter-1]
            table.cell(
                counter, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(
                counter, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(
                counter, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.rows[counter].height = Cm(0.8)
            counter = counter + 1

    i = 0
    for e in document.paragraphs:
        #print(e.text + str(i))
        i = i + 1
    currentDate = datetime.now().date()
    #document.paragraphs[12].add_run(
    #    " " + str(currentDate.day) + "." + str(currentDate.month) + "." + str(currentDate.year))

    table = document.tables[1]
    table.cell(0, 0).paragraphs[0].add_run("\n" + reverzJson["uporabnik"])
    table.cell(0, 1).paragraphs[0].add_run("\n" + reverzJson["EZSO"])
    table.cell(0, 0).paragraphs[0].runs[0].font.bold = True

    table = document.tables[0]
    table.cell(0, 0).paragraphs[0].add_run(" " + reverzJson["unit"])
    table.cell(1, 0).paragraphs[0].add_run(" " + reverzJson["STRM"])

    user = reverzJson["me"]
    useSignature = reverzJson["useSignature"]
    employee = reverzJson["uporabnik"]
    givenLocation = reverzJson["izdanaLokacija"]
    receivedLocation = reverzJson["vrnjenaLokacija"]
    reverzCounter = reverzJson["reverzCounter"]
    signatureOffset = reverzJson["signatureOffset"]
    star = reverzJson["star"]
    print(f"tuka: {reverzJson['vrnjenaLokacija']}")
    if star == "1":
        star = "Oznaka * pri inventarni številki predstavlja opremo, ki je bila izdana iz skladišča"
    else:
        star = ""

    document.save(f"{sessionUser}/tmp2.docx")
    bash_given = givenLocation.replace("/", "\/").replace("#", "\#").replace("!", "\!")
    bash_taken = receivedLocation.replace("/", "\/").replace("#", "\#").replace("!", "\!")
    print(f"TUDIIII {bash_taken}")
    os.system(
        f"./reverzComplete.sh '{sessionUser}' '{user}' '{employee}' '{bash_given}' '{bash_taken}' '{reverzCounter}' '{star}'")
    os.system(f"unoconv -f pdf {sessionUser}/tmp2.docx")
    
    useSignature = int(useSignature)
    if useSignature == 1:
        print(signatureOffset)
        in_pdf_file = f'{sessionUser}/tmp2.pdf'
        out_pdf_file = f'{sessionUser}/tmp.pdf'
        img_file = f'{sessionUser}/signature.png'

        packet = io.BytesIO()
        can = canvas.Canvas(packet)
        x_start = 70
        y_start = 60 - int(signatureOffset) *0.44 + 3
        can.drawImage(img_file, x_start, y_start, width=120, height=45, preserveAspectRatio=True, mask='auto')
        can.save()

        #move to the beginning of the StringIO buffer
        packet.seek(0)

        new_pdf = PdfFileReader(packet)

        # read the existing PDF
        existing_pdf = PdfFileReader(open(in_pdf_file, "rb"))
        output = PdfFileWriter()

        for i in range(len(existing_pdf.pages)):
            page = existing_pdf.getPage(i)
            page.mergePage(new_pdf.getPage(i))
            output.addPage(page)

        outputStream = open(out_pdf_file, "wb")
        output.write(outputStream)
        outputStream.close()

    else:
        os.system(f"cp {sessionUser}/tmp2.pdf {sessionUser}/tmp.pdf")

        #document._body._body.xml = xmlBody
        #print("")